import asyncio
from asyncio import subprocess
from io import BytesIO
import os
import tempfile
import threading
from typing import Dict, List, Optional, Any
import uuid
from fastapi import File, UploadFile
import pandas as pd
from docx import Document as DocxDocument 
from langchain_community.document_loaders import PyPDFLoader
from langchain.embeddings.base import Embeddings 
from src.backend.utils.async_runner import AsyncRunner

asyncrunner = AsyncRunner()

# async def upload_files(user_id: str, file: UploadFile):
    
#     return await parse_docs(file, user_id, file_id)

async def extract_text_from_file(file: UploadFile) -> str:
    """Extracts text content from an uploaded file based on its type."""
    filename = file.filename
    content = await file.read()  # Read the content asynchronously
    file_extension = filename.split(".")[-1].lower()

    try:
        if file_extension == "pdf":
            # Correct: synchronous function to be run in a thread
            def _extract():
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                    temp_pdf.write(content)
                    temp_pdf_path = temp_pdf.name
                try:
                    loader = PyPDFLoader(temp_pdf_path)
                    documents = loader.load()  # Blocking call
                    extracted_text_pages = [doc.page_content.strip() for doc in documents]
                    return "\n\n".join(extracted_text_pages)
                finally:
                    os.remove(temp_pdf_path)

            return await asyncio.to_thread(_extract)

        elif file_extension == "txt":
            return content.decode("utf-8")

        elif file_extension in ["xlsx", "xls", "csv"]:
            def _extract_excel():
                if file_extension == "csv":
                    df = pd.read_csv(BytesIO(content))
                else:
                    df = pd.read_excel(
                        BytesIO(content), engine="xlrd" if file_extension == "xls" else "openpyxl"
                    )
                return df.astype(str).apply(lambda row: "|".join(row), axis=1).str.cat(sep='\n')

            return await asyncio.to_thread(_extract_excel)

        elif file_extension == "docx":
            def _extract_docx():
                doc = DocxDocument(BytesIO(content))
                return "\n".join([para.text for para in doc.paragraphs])

            return await asyncio.to_thread(_extract_docx)

        elif file_extension == "doc":
            def _extract_doc():
                with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_doc:
                    temp_doc.write(content)
                    temp_doc_path = temp_doc.name
                try:
                    result = subprocess.run(["antiword", temp_doc_path], text=True, capture_output=True, check=True)
                    return result.stdout
                finally:
                    os.remove(temp_doc_path)

            return await asyncio.to_thread(_extract_doc)

        elif file_extension == "md":
            return content.decode("utf-8")

        else:
            return f"Unsupported file format: .{file_extension}"

    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        return f"Error extracting text from {filename}: {e}"


async def upload_files(user_id, file):
    """
    Upload and process a file, storing its embeddings in Qdrant.
    
    Args:
        file: The uploaded file object
        user_id: The user ID uploading the file
    
    Returns:
        str: The file_id if successful
    
    Raises:
        DocumentAlreadyExistsError: If the document already exists
        ValueError: If the file format is unsupported
        Exception: For other processing errors
    """
    try:
        file_id = str(uuid.uuid4())
        print(f"Processing file: {file.filename}")
        
        # Process file based on extension
        if file.filename.lower().endswith(".pdf"):
            content = await process_pdf(file)
        elif file.filename.lower().endswith(".txt"):
            content = await read_txt_file(file)
        elif file.filename.lower().endswith((".xlsx", ".xls", ".csv")):
            content = await process_excel(file)
        elif file.filename.lower().endswith(".doc"):
            content = await process_doc(file)
        elif file.filename.lower().endswith(".docx"):
            content = await process_docx(file)
        elif file.filename.lower().endswith(".md"):
            content = await process_md(file)
        else:
            raise ValueError(f"Unsupported file format: {file.filename}")

        # Store document embeddings
        # result = await qdrant.store_document_embeddings(
        #     file_id=file_id,
        #     filename=file.filename,
        #     text_content=content,
        #     user_id=user_id
        # )
        
        # # Handle the result based on status
        # if result["status"] == "success":
        #     print(f"Successfully uploaded file: {file.filename} with ID: {result.get('file_id')}")
        #     return result.get("file_id")
        # else:
        #     raise Exception(f"Failed to store document: {result.get('message', 'Unknown error')}")
        return ""
        
    except Exception as e:
        raise Exception(f"Unexpected status from storage {e}")
        
async def process_docx(self, file: UploadFile):
    try:
        from docx import Document
        doc = Document(BytesIO(file.file.read()))
        text = "\n".join([para.text for para in doc.paragraphs])
        print(f"\n--- Extracted DOCX Text ---\n{text}")
        return text
    except Exception as e:
        print(f"Error processing DOCX: {e}")
        return ""

async def process_md(self, file: UploadFile):
    try:
        content = file.file.read().decode("utf-8")
        print(f"\n--- Extracted Markdown Text ---\n{content}\n")
        return content
    except Exception as e:
        print(f"Error processing Markdown file: {e}")
        return ""

async def process_doc(file: UploadFile):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_doc:
            temp_doc.write(file.file.read())
            temp_doc_path = temp_doc.name

        print(f"Temporary file created: {temp_doc_path}")
        result = subprocess.run(["antiword", temp_doc_path], text=True, capture_output=True)
        text = result.stdout
        os.remove(temp_doc_path)
        print(f"\n--- Extracted DOC Text ---\n{text}")
        return text
    except Exception as e:
        print(f"Error processing DOC: {e}")
        return ""

async def read_txt_file(file: UploadFile):
    content = file.file.read().decode("utf-8")
    print(f"\n--- {file.filename} (TXT File) ---\n{content}\n")
    return content

async def process_pdf(file: UploadFile) -> Dict[str, Any]: 
    """Extracts text from a PDF, calls the LLM validator, and returns its JSON response."""
    # 1. Dump to temp file
    try:
        file.file.seek(0)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(file.file.read())
            temp_pdf_path = temp_pdf.name
        # 2. Load raw text per page
        loader = PyPDFLoader(temp_pdf_path)
        documents = loader.load()
        extracted_text = [doc.page_content.strip() for doc in documents]
        combined = "\n\n".join(extracted_text)
        return combined
    except Exception as e:
        print(f"PDF load failed: {e}")
        return None
    finally:
        # cleanup temp file
        os.remove(temp_pdf_path)


def clean_excel_content(pipe_separated_list):
    """
    Clean the pipe-separated Excel data to extract meaningful content
    """
    if not pipe_separated_list:
        return ""
    
    cleaned_lines = []
    
    for line in pipe_separated_list:
        if not line or not isinstance(line, str):
            continue

        parts = [
            part.strip() 
            for part in line.split('|') 
            if part.strip() and part.strip().lower() not in ['nan', 'none', '']
        ]

        if parts:
            cleaned_line = ' '.join(parts)
            if cleaned_line and len(cleaned_line.strip()) > 0:
                cleaned_lines.append(cleaned_line)
    

    final_content = '\n'.join(cleaned_lines)

    if not final_content.strip():
        return "Document processed but contains no readable text content."
    
    return final_content

async def process_excel(file: UploadFile):
    try:
        contents = file.file.read()
        file_extension = file.filename.split(".")[-1]
        
        if file_extension == "csv":
            df = pd.read_csv(BytesIO(contents))
        elif file_extension in ["xls", "xlsx"]:
            df = pd.read_excel(
                BytesIO(contents), engine="xlrd" if file_extension == "xls" else "openpyxl"
            )
        else:
            raise ValueError("Unsupported file format!")

        df = (
            df.astype(str)
            .apply(lambda row: "|".join(row), axis=1)
            .to_frame(name="Pipe_Separated")
        )
        
        # Create output file
        output = BytesIO()
        if file_extension == "csv":
            df.to_csv(output, sep="|", index=False, header=False)
        else:
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                df.to_excel(writer, index=False, sheet_name="Sheet1")
        output.seek(0)

        pipe_separated_data = df['Pipe_Separated'].tolist()
        cleaned_content = clean_excel_content(pipe_separated_data)
        
        return cleaned_content
        
    except Exception as e:
        print(f"Error processing file: {e}")
        return ""
