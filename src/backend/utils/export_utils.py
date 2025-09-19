import asyncio
import logging
from pathlib import Path
import io 
import os  
import functools
# PDF Generation
import pdfkit

# Markdown Parsing
from markdown_it import MarkdownIt
from markdown_it.rules_inline import StateInline
from markdown_it.token import Token

# mdit_py_plugins:
from mdit_py_plugins.amsmath import amsmath_plugin
from mdit_py_plugins.deflist import deflist_plugin
from mdit_py_plugins.footnote import footnote_plugin
from mdit_py_plugins.front_matter import front_matter_plugin
from mdit_py_plugins.texmath import texmath_plugin
from mdit_py_plugins.tasklists import tasklists_plugin

# Syntax Highlighting with Pygments
from pygments import highlight as pygments_highlight
from pygments.lexers import get_lexer_by_name, guess_lexer
from pygments.formatters import HtmlFormatter
from pygments.util import ClassNotFound as PygmentsClassNotFound

# DOCX Generation
# Factory for creating new documents
from docx import Document as DocumentFactory
# Actual Document class for isinstance checks
from docx.document import Document as DocxDocument

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn  
from docx.oxml import OxmlElement  
from bs4 import BeautifulSoup  
import httpx
import re


def slugify(value: str, max_length: int = 40) -> str:
    s = re.sub(r'[^a-z0-9]+', '-', value.lower()).strip('-')
    return s[:max_length] or "export"


# # --- Basic Logging Setup ---
# logging.basicConfig(
#     level=logging.INFO,
#     format='%(asctime)s - %(levelname)s - %(module)s - %(message)s'
# )


# --- PDF: Core Configuration ---
WKHTMLTOPDF_PATH_PDF = '/usr/bin/wkhtmltopdf'# Example: '/usr/local/bin/wkhtmltopdf'
PDFKIT_CONFIG = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH_PDF) if os.path.exists(WKHTMLTOPDF_PATH_PDF) else None

CURRENT_DIR = Path(__file__).resolve().parent
KATEX_CSS_PATH = "static/katex.min.css"
PYGMENTS_CSS_PATH = "static/pygments_style.css"


def load_css_file(file_path: Path, description: str) -> str:
    try:
        if file_path.exists():
            return file_path.read_text(encoding='utf-8')
        logging.warning(f"{description} CSS file NOT FOUND at: {file_path}")
        return f"/* {description} CSS not found at {file_path}. Please ensure the file exists. */"
    except Exception as e:
        logging.error(f"Error reading {description} CSS from {file_path}: {e}")
        return f"/* Error reading {description} CSS from {file_path}. Check permissions and file integrity. */"


KATEX_CSS = load_css_file(KATEX_CSS_PATH, "KaTeX for PDF")
PYGMENTS_CSS = load_css_file(PYGMENTS_CSS_PATH, "Pygments for PDF")

CUSTOM_CSS_PDF = """
body { font-family: 'Inter', Arial, sans-serif; line-height: 1.7; margin: 25mm; color: #333333; max-width: 210mm; box-sizing: border-box; }
h1, h2, h3, h4, h5, h6 { color: #111111; margin-top: 1.5em; margin-bottom: 0.6em; line-height: 1.3; }
h1 { font-size: 24pt; border-bottom: 1px solid #dddddd; padding-bottom: 0.2em;}
h2 { font-size: 18pt; border-bottom: 1px solid #eeeeee; padding-bottom: 0.2em;}
h3 { font-size: 14pt; } h4 { font-size: 12pt; font-style: italic;}
p { margin-top: 0.5em; margin-bottom: 0.5em; font-size: 10pt; }
a { color: #0066cc; text-decoration: none; } a:hover { text-decoration: underline; }
code:not(pre > code) { background-color: #f0f0f0; padding: .1em .3em; border-radius: 3px; font-family: 'Courier New', Courier, monospace; font-size: 0.9em; }
.highlight pre { background-color: #f8f9fa; border: 1px solid #dee2e6; border-radius: 4px; padding: 1em; overflow-x: auto; font-size: 9pt; line-height: 1.45; }
table { border-collapse: collapse; width: 100%; margin-bottom: 1em; font-size: 9pt;}
th, td { border: 1px solid #cccccc; padding: 8px; text-align: left; }
thead th { background-color: #f2f2f2; font-weight: bold; }
blockquote { border-left: 4px solid #cccccc; padding-left: 1em; margin-left: 0; color: #555555; font-style: italic; }
img { max-width: 100%; height: auto; margin-top: 0.5em; margin-bottom: 0.5em; border-radius: 4px; }
ul, ol { padding-left: 20px; }
.katex-display { margin: 1em 0; overflow-x: auto; text-align: center; }
.katex-display > .katex { font-size: 1.1em; }
"""
COMBINED_CSS_PDF = f"<style>\n{KATEX_CSS}\n{PYGMENTS_CSS}\n{CUSTOM_CSS_PDF}\n</style>"


# --- Custom Mark Plugin for ==highlight== ---
def mark_plugin(md: MarkdownIt):
    def tokenize_mark(state: StateInline, silent: bool):
        src = state.src
        pos = state.pos
        if src[pos:pos+2] != '==' or pos + 4 > len(src):
            return False
        end = src.find('==', pos+2)
        if end == -1:
            return False
        if not silent:
            token_open = Token('mark_open', 'mark', 1)
            state.push(token_open)
            token_text = Token('text', '', 0)
            token_text.content = src[pos+2:end]
            state.push(token_text)
            token_close = Token('mark_close', 'mark', -1)
            state.push(token_close)
        state.pos = end + 2
        return True

    md.inline.ruler.after('strikethrough', 'mark', tokenize_mark)


# --- Syntax Highlighting Function for Markdown-it-py ---
def highlight_code_with_pygments(code_str: str, language: str, attrs: str) -> str:
    try:
        lexer = get_lexer_by_name(language) if language else guess_lexer(code_str)
    except PygmentsClassNotFound:
        logging.warning(f"Pygments lexer not found for '{language}'. Using 'text'.")
        lexer = get_lexer_by_name("text")
    formatter = HtmlFormatter(cssclass="highlight", linenos=False, nowrap=False)
    return pygments_highlight(code_str, lexer, formatter)


# --- Markdown Parser Setup (used by both PDF and DOCX) ---
md_parser = (
    MarkdownIt(
        "commonmark",
        {"html": True, "linkify": True, "typographer": True, "highlight": highlight_code_with_pygments},
    )
    .enable(["table", "strikethrough"])
    .use(front_matter_plugin)
    .use(footnote_plugin)
    .use(amsmath_plugin)
    .use(texmath_plugin, delimiters='dollars', macros={r"\RR": r"\mathbb{R}"})
    .use(deflist_plugin)
    .use(tasklists_plugin, enabled=True)
    .use(mark_plugin)
)

# --- HTML Preprocessing (shared logic) ---
def preprocess_html_content(html_content: str, base_url: str = None) -> str:
    processed_html = html_content
    if base_url:
        base_url_with_slash = base_url if base_url.endswith('/') else base_url + '/'
        for old, new in {
            'src="public/': f'src="{base_url_with_slash}public/',
            "src='/public/": f'src="{base_url_with_slash}public/',
            'href="public/': f'href="{base_url_with_slash}public/',
            "href='/public/": f'href="{base_url_with_slash}public/',
        }.items():
            processed_html = processed_html.replace(old, new)
    return processed_html

# --- PDF Export Function ---
async def markdown_to_pdf(markdown_content: str, pdf_file_path: str, base_url: str = None):
    try:
        if not markdown_content:
            logging.warning("Markdown content for PDF is empty.")
        html_body = md_parser.render(markdown_content)
        html_body_processed = preprocess_html_content(html_body, base_url)
        html_full_document = f"<!DOCTYPE html><html lang='en'><head><meta charset='utf-8'><title>Generated PDF</title>{COMBINED_CSS_PDF}</head><body>{html_body_processed}</body></html>"
        pdf_options = {
            'encoding': "UTF-8", 'enable-local-file-access': None,
            'margin-top': '0.75in', 'margin-right': '0.75in',
            'margin-bottom': '0.75in', 'margin-left': '0.75in',
            'page-size': 'A4', 'quiet': None,
        }
        pdf_options_cleaned = {k: v for k, v in pdf_options.items() if v is not None}
        loop = asyncio.get_event_loop()
        func = functools.partial(
            pdfkit.from_string,
            html_full_document,
            pdf_file_path,
            options=pdf_options_cleaned,
            configuration=PDFKIT_CONFIG
        )
        success = await loop.run_in_executor(None, func)
        if success:
            logging.info(f"Successfully created PDF: {pdf_file_path}")
        else:
            logging.error(f"PDF generation reported failure for: {pdf_file_path}")
    except FileNotFoundError as e:
        logging.error(f"WKHTMLTOPDF for PDF ERROR: Not found. {e}")
    except IOError as e:
        logging.error(f"WKHTMLTOPDF/IO ERROR for PDF: {e}")
    except Exception as e:
        logging.error(f"Unexpected error during PDF conversion: {e}", exc_info=True)


# --- DOCX Export Helpers and Function ---
# (Same as original implementation)

# --- DOCX Export Function ---
async def _add_html_element_to_docx(element, doc_or_para, base_url):
    # Use DocxDocument for isinstance checks
    is_adding_to_paragraph = not isinstance(doc_or_para, DocxDocument)
    doc = doc_or_para if isinstance(doc_or_para, DocxDocument) else doc_or_para.part.document
    current_paragraph = doc_or_para if is_adding_to_paragraph else None

    if element.name is None:  # NavigableString
        if element.strip():
            if current_paragraph is None:
                if doc.paragraphs and (not doc.paragraphs[-1].runs or not doc.paragraphs[-1].text.strip()):
                    current_paragraph = doc.paragraphs[-1]
                elif doc.paragraphs and doc.paragraphs[-1].runs and not doc.paragraphs[-1].text.endswith(("\n", " ", "\t", u'\xa0')):
                    current_paragraph = doc.paragraphs[-1]
                else:
                    current_paragraph = doc.add_paragraph()
            run = current_paragraph.add_run(str(element))
            parent = element.parent
            if parent:
                if parent.name in ['strong', 'b']:
                    run.bold = True
                if parent.name in ['em', 'i']:
                    run.italic = True
                if parent.name == 'code' and (parent.parent is None or parent.parent.name != 'pre'):
                    run.font.name = 'Courier New'
                    try:
                        rpr = run._r.get_or_add_rPr()
                        cs_font = OxmlElement('w:font')
                        cs_font.set(qn('w:cs'), 'Courier New')
                        rpr.append(cs_font)
                    except Exception as e_font:
                        logging.debug(f"Could not set cs font for inline code: {e_font}")
        return

    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(element.name[1])
        heading_p = doc.add_heading(level=level)
        for child in element.children:
            await _add_html_element_to_docx(child, heading_p, base_url)
    elif element.name == 'p':
        para = doc.add_paragraph()
        for child in element.children:
            await _add_html_element_to_docx(child, para, base_url)
    elif element.name == 'hr':
        doc.add_paragraph("_________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif element.name in ['ul', 'ol']:
        for li in element.find_all('li', recursive=False):
            style = 'ListBullet' if element.name == 'ul' else 'ListNumber'
            list_item_p = doc.add_paragraph(style=style)
            for child in li.children:
                await _add_html_element_to_docx(child, list_item_p, base_url)
    elif element.name == 'pre' or (element.name == 'div' and "highlight" in element.get("class", [])):
        actual_pre = element if element.name == 'pre' else element.find('pre')
        code_block_text = actual_pre.get_text() if actual_pre else element.get_text()
        code_p = doc.add_paragraph(style='NoSpacing')
        run = code_p.add_run(code_block_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        try:
            rpr = run._r.get_or_add_rPr()
            cs_font = OxmlElement('w:font')
            cs_font.set(qn('w:cs'), 'Courier New')
            rpr.append(cs_font)
        except Exception as e_font:
            logging.debug(f"Could not set cs font for code block: {e_font}")
    elif element.name == 'img':
        src = element.get('src')
        alt = element.get('alt', 'image')
        if src:
            try:
                is_url = src.startswith(('http://', 'https://'))
                is_abs_path = Path(src).is_absolute()
                if not (is_url or is_abs_path):
                    if base_url and base_url.startswith('file:///'):
                        base_path_str = base_url[len('file:///'):].lstrip('/')
                        if os.name == 'nt' and len(base_path_str) > 1 and base_path_str[0] == '/' and base_path_str[2] == ':':
                            base_path_str = base_path_str[1:]
                        src = str((Path(base_path_str) / src.lstrip('/')).resolve())
                        is_abs_path = True
                if is_url:
                    async with httpx.AsyncClient() as client:
                        response = await client.get(src, timeout=10)
                        response.raise_for_status()
                    doc.add_picture(io.BytesIO(response.content), width=Inches(4.5))
                    logging.info(f"Added image to DOCX from URL: {src}")
                elif is_abs_path and Path(src).is_file():
                    doc.add_picture(src, width=Inches(4.5))
                    logging.info(f"Added image to DOCX from local file: {src}")
                else:
                    logging.warning(f"Image src '{src}' for DOCX could not be resolved. Skipping.")
                if alt:
                    cap_p = doc.add_paragraph(style='Caption')
                    cap_p.add_run(alt).italic = True
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logging.error(f"Error adding image {src} to DOCX: {e}", exc_info=True)
    elif element.name == 'table':
        rows_data, header_data = [], []
        thead = element.find('thead')
        if thead and thead.find('tr'):
            header_data = [th.get_text(strip=True) for th in thead.find('tr').find_all(['th', 'td'])]
        tbody = element.find('tbody') if element.find('tbody') else element
        for tr in tbody.find_all('tr', recursive=False):
            if thead and tr.parent.name == 'thead':
                continue
            rows_data.append([td.get_text(strip=True) for td in tr.find_all(['td', 'th'])])
        if header_data or rows_data:
            num_cols = len(header_data) if header_data else (len(rows_data[0]) if rows_data else 0)
            if num_cols > 0:
                table_docx = doc.add_table(rows=1 if header_data else 0, cols=num_cols)
                table_docx.style = 'TableGrid'
                if header_data:
                    hdr_cells = table_docx.rows[0].cells
                    for i, col_text in enumerate(header_data):
                        if i < num_cols:
                            hdr_cells[i].text = col_text
                            for paragraph in hdr_cells[i].paragraphs:
                                for run_item in paragraph.runs:
                                    run_item.bold = True
                for row_dt in rows_data:
                    row_cells_docx = table_docx.add_row().cells
                    for i, cell_text in enumerate(row_dt):
                        if i < num_cols:
                            row_cells_docx[i].text = cell_text
    elif element.name == 'blockquote':
        quote_p = doc.add_paragraph()
        quote_p.paragraph_format.left_indent = Inches(0.5)
        for child in element.children:
            await _add_html_element_to_docx(child, quote_p, base_url)
        doc.add_paragraph()
    elif element.name in ['strong', 'b', 'em', 'i', 'span', 'a', 'code', 'div', 'figure', 'figcaption', 'body', 'html']:
        target_para_or_doc = current_paragraph if current_paragraph else doc
        for child in element.children:
            await _add_html_element_to_docx(child, target_para_or_doc, base_url)
    elif element.name in ['script', 'style', 'meta', 'link', 'head', 'title', 'annotation']:
        pass
    else:
        if not is_adding_to_paragraph:
            unknown_p = doc.add_paragraph()
            text_content = element.get_text(separator=' ', strip=True)
            if text_content:
                unknown_p.add_run(text_content)
            logging.debug(f"Unhandled HTML element <{element.name}> in DOCX, added as text.")

async def _parse_html_to_docx_doc(html_content: str, doc: DocxDocument, base_url: str = None):
    """Parses HTML content and populates a python-docx Document object."""
    soup = BeautifulSoup(html_content, 'lxml')
    for math_element in soup.find_all(class_=["math-display", "math-inline"]):
        latex_source_tag = math_element.find('annotation', encoding='application/x-tex')
        if latex_source_tag:
            latex_source = latex_source_tag.get_text(strip=True)
            p_math = doc.add_paragraph()
            prefix = "Display Math: " if "math-display" in math_element.get('class', []) else "Inline Math: "
            p_math.add_run(prefix).italic = True
            math_run = p_math.add_run(f" {latex_source} ")
            math_run.font.name = 'Courier New'
            math_run.font.size = Pt(9)
        math_element.decompose()
    body_content = soup.body if soup.body else soup
    if body_content:
        for element in body_content.children:
            await _add_html_element_to_docx(element, doc, base_url)

async def markdown_to_docx(markdown_content: str, docx_file_path: str, base_url: str = None):
    """Converts a Markdown string to a DOCX file."""
    try:
        logging.info(f"Starting DOCX conversion for: {docx_file_path}")
        if not markdown_content:
            logging.warning("Markdown content for DOCX is empty.")
        html_body = md_parser.render(markdown_content)
        html_body_processed = preprocess_html_content(html_body, base_url)

        doc = DocumentFactory()

        await _parse_html_to_docx_doc(html_body_processed, doc, base_url)
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, doc.save, docx_file_path)
        logging.info(f"Successfully created DOCX: {docx_file_path}")
    except Exception as e:
        logging.error(f"Unexpected error during DOCX conversion: {e}", exc_info=True)
